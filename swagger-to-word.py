import requests
import json
import docx
swagger_url = "http://IP:Port/v3/api-docs"  # Adjust this if necessary
try:
    from docx import Document
except ImportError:
    print("Error: The 'python-docx' module is not installed. Please install it using 'pip install python-docx'")
    exit(1)

def fetch_swagger_json(url):
    response = requests.get(url)
    if response.status_code == 200:
        return response.json()
    else:
        print(f"Failed to fetch Swagger JSON: {response.status_code}")
        return None

def extract_endpoints(swagger_json):
    endpoints = {}
    paths = swagger_json.get("paths", {})

    for path, methods in paths.items():
        for method, details in methods.items():
            tag = details.get("tags", ["Uncategorized"])[0]
            if tag not in endpoints:
                endpoints[tag] = []
            endpoints[tag].append(f"{method.upper()} {path}")

    return endpoints

def generate_word_document(endpoints, filename="swagger_endpoints.docx"):
    doc = Document()
    doc.add_heading("API Endpoints", level=1)

    for category, paths in endpoints.items():
        doc.add_heading(category, level=2)
        for endpoint in paths:
            doc.add_paragraph(endpoint)

    doc.save(filename)
    print(f"Word document saved as {filename}")

def main():
    swagger_json = fetch_swagger_json(swagger_url)

    if swagger_json:
        endpoints = extract_endpoints(swagger_json)
        generate_word_document(endpoints)

if __name__ == "__main__":
    main()
